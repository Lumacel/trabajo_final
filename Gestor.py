from tkinter import *
from tkinter import ttk, messagebox,filedialog
from docx import Document
from docx.shared import Pt, Cm
from datetime import datetime 
import csv
import os

class App:
	def __init__(self):
		self.afiliados = []
		self.nuevo_afiliado=[]
		self.afiliado_target = []
		self.obras_sociales = []
		self.vista_obras_sociales= []
		self.current_obrasocial="TODAS"
		
		self.root = Tk()  
		self.root.title("GESTOR") 
		self.root.resizable(0,0)
		self.centrar_ventana()

		self.cargar_obras_sociales()

		# -- agrega estilo 
		self.style = ttk.Style()
		self.style.theme_use('winnative') 

		# -- label para información
		self.text_info = StringVar()
		self.text_info_frame1= StringVar()

		self.lbl_info = Label(self.root, textvariable = self.text_info , width=114, relief="ridge",bg="#e1f3fd",fg="blue", bd=4)
		self.lbl_info.place(x=47,y=240)
		
		# -- Nueva ventana dentro de la raiz(root)
		self.frame1 = Frame(self.root,bg="lightgrey",relief="ridge",bd=4)
		self.frame1.place(x=48,y=282)
		self.frame1.configure(width=808,height=121)

		self.separador= ttk.Separator(self.frame1, orient="vertical")
		self.separador.place(relx=0.33, rely=0, relwidth=0.005, relheight=1) 

		self.separador= ttk.Separator(self.frame1, orient="vertical")
		self.separador.place(relx=0.66, rely=0, relwidth=0.005, relheight=1)

		self.btn_agregar= Button(self.frame1,text="AGREGAR", width=29, bd=3, command= self.modo_agregar_afiliado)
		self.btn_agregar.place(x=26,y=10) 

		self.btn_editar= Button(self.frame1,text="EDITAR", width=29, bd=3, command= self.modo_editar_afiliado)
		self.btn_editar.place(x=26,y=45)

		self.btn_eliminar= Button(self.frame1,text="ELIMINAR", width=29, bd=3, command= self.modo_eliminar_afiliado)
		self.btn_eliminar.place(x=26,y=80)

		self.lbl_obra_social= Label(self.frame1,text="OBRA SOCIAL", relief= "ridge",width=14)
		self.lbl_obra_social.place(x=294,y=10)

		self.entry_obra_social = ttk.Combobox(self.frame1,values=self.vista_obras_sociales,state="readonly",width=12)
		self.entry_obra_social.place(x=407,y=10)
		self.entry_obra_social.current(0)

		self.lbl_info_frame1 = Label(self.frame1,textvariable = self.text_info_frame1 , relief= "ridge",width=29, bg="#e1f3fd", fg="black")
		self.lbl_info_frame1.place(x=294,y=45)

		#self.btn_actualizar= Button(self.frame1,text="ACTUALIZAR TABLA",width=29, bd=3, command= self.actualizar_tabla)
		#self.btn_actualizar.place(x=294,y=80)

		self.btn_crear_doc= Button(self.frame1,text="GENERAR DOCUMENTO",width=29, bd=3, command= self.generar_documento)
		self.btn_crear_doc.place(x=562,y=10)

		self.btn_carpeta_doc= Button(self.frame1,text="EXAMINAR DOCUMENTOS",width=29, bd=3, command= self.abrir_carpeta_docs)
		self.btn_carpeta_doc.place(x=562,y=45)

		self.root.bind("<<ComboboxSelected>>", lambda x : self.actualizar_tabla() )

		self.crear_tabla()
		self.inicializar_tabla()

	def centrar_ventana(self):
		ancho_ventana = 900
		alto_ventana = 420
		ancho_pantalla= self.root.winfo_screenwidth()
		alto_pantalla=self.root.winfo_screenheight()
		x_ventana = ancho_pantalla // 2 - ancho_ventana // 2
		y_ventana = alto_pantalla // 2 - alto_ventana // 2
		posicion = str(ancho_ventana) + "x" + str(alto_ventana) + "+" + str(x_ventana) + "+" + str(y_ventana)
		self.root.geometry(posicion)

	def inicializar_tabla(self):
		self.crear_lista_afiliados() 
		self.ordenar_lista_afiliados()
		self.cargar_tabla(self.get_afiliados_sort(),self.get_obrasocial())
		self.info_num_afiliados()

	def crear_tabla(self,filas=10): # --- da formato a la tabla (treeview)
		self.tabla = ttk.Treeview(self.root, height=filas,columns=("APELLIDO","NOMBRE","DNI/AFILIADO","TELEFONO","OBRA SOCIAL"))
		self.tabla.pack(expand=False)
		
		# --- agrega color al heading
		self.style.configure('Treeview.Heading', background="#e1f3fd")
	
		# --- barra scroll
		self.scroll = Scrollbar(self.root, orient="vertical", command=self.tabla.yview)
		self.scroll.place(x=860, y=10, height=220)
		self.tabla.configure(yscrollcommand=self.scroll.set)

		# --- formato a las columnas
		self.tabla.column("#0", width=0, stretch=NO , minwidth=100)
		self.tabla.column("APELLIDO", anchor=W, width=200, minwidth = 200)
		self.tabla.column("NOMBRE", anchor=W, width=200, minwidth = 200)
		self.tabla.column("DNI/AFILIADO", anchor=E, width=150, minwidth = 150)
		self.tabla.column("TELEFONO", anchor=E, width=150, minwidth = 150)
		self.tabla.column("OBRA SOCIAL", anchor=CENTER, width=100,minwidth = 100)

		# --- indicar cabecera
		self.tabla.heading("#0", text="", anchor=CENTER)
		self.tabla.heading("#1", text="APELLIDO", anchor=CENTER)
		self.tabla.heading("#2", text="NOMBRE", anchor=CENTER)
		self.tabla.heading("#3", text="DNI/AFILIADO", anchor=CENTER)
		self.tabla.heading("#4", text="TELEFONO", anchor=CENTER)
		self.tabla.heading("#5", text="OBRA SOCIAL", anchor=CENTER)

	def abrir_ventana(self): # --- configura ventana para entrada de datos
		self.nuevo_afiliado=[]
		self.top_level = Toplevel()
		self.top_level.resizable(0,0)
		self.top_level.title("DATOS DEL AFILIADO")
		self.top_level.geometry("370x210+200+200")

		self.apellido = StringVar()
		self.nombre = StringVar() 
		self.dni_afiliado = StringVar()
		self.telefono = StringVar()
		self.obrasocial= StringVar()
		
		self.lbl_apellido = Label(self.top_level, text = "APELLIDO", width=20, relief="ridge", bg="#e1f3fd")
		self.lbl_apellido.place(x=10,y=10)
		self.entry_apellido = Entry(self.top_level,textvariable = self.apellido,width=30)
		self.entry_apellido.place(x=170,y=10)

		self.lbl_nombre = Label(self.top_level, text = "NOMBRE", width=20, relief="ridge", bg="#e1f3fd")
		self.lbl_nombre.place(x=10,y=40)
		self.entry_nombre = Entry(self.top_level,textvariable = self.nombre,width=30)
		self.entry_nombre.place(x=170,y=40) 

		self.lbl_dni_afiliado = Label(self.top_level, text = "DNI/AFILIADO", width=20, relief="ridge", bg="#e1f3fd")
		self.lbl_dni_afiliado.place(x=10,y=70)
		self.lbl_dni_afiliado= Entry(self.top_level,textvariable = self.dni_afiliado,width=30)
		self.lbl_dni_afiliado.place(x=170,y=70)

		self.lbl_telefono = Label(self.top_level, text = "TELEFONO", width=20, relief="ridge", bg="#e1f3fd")
		self.lbl_telefono.place(x=10,y=100)
		self.entry_telefono = Entry(self.top_level,textvariable = self.telefono,width=30)
		self.entry_telefono.place(x=170,y=100)

		self.lbl_obra_social = Label(self.top_level, text = "OBRA SOCIAL", width=20, relief="ridge", bg="#e1f3fd")
		self.lbl_obra_social.place(x=10,y=130)
		self.entry_obrasocial_toplevel = ttk.Combobox(self.top_level,textvariable = self.obrasocial,values=self.obras_sociales,state="readonly",width=6)
		self.entry_obrasocial_toplevel.place(x=170,y=130)

		self.btn_ingresar= Button(self.top_level,text="GRABAR DATOS",bd=3, width=15, command= self.grabar_afiliado)
		self.btn_ingresar.place(x=10,y=175)

		self.btn_salir= Button(self.top_level,text="SALIR", bd=3, width=15, command= self.salir_ventana)
		self.btn_salir.place(x=245,y=175)
		
		self.entry_apellido.focus()

		self.top_level.grab_set() # --- inhabilita controles ventana principal

	def cargar_obras_sociales(self,archivo = "obras_sociales/obras_sociales.csv"): # --- levanta datos de obras sociales de archivo csv
		try:
			with open(archivo, 'r', encoding='latin1') as ob_soc:
				csvreader = csv.reader(ob_soc) 
				for linea in csvreader:
					if linea == []:continue
					self.obras_sociales.append(linea[0].upper())
				self.vista_obras_sociales.append("TODAS")
				for i in range(len(self.obras_sociales)):
					self.vista_obras_sociales.append(self.obras_sociales[i])
		except Exception as e:
			messagebox.showerror(message=e, title="ERROR!!!")

	def crear_lista_afiliados(self,archivo="afiliados/afiliados.csv"): # --- levanta datos de archivo .csv (crea lista de lista de datos)		
		try:
			with open(archivo, 'r', encoding='latin1') as datos:
				csvreader = csv.reader(datos)
				items = next(csvreader) 
				for linea in csvreader:			
					if linea == []:continue
					self.afiliados.append(linea)	
				self.items = items
		except Exception as e:
			messagebox.showerror(message=e, title="ERROR!!!")

	def cargar_tabla(self,lista_afiliados,obrasocial="TODAS"): # --- agrega datos a la tabla (treeview)
		for afiliado in lista_afiliados: 
			if obrasocial != "TODAS":
				if afiliado[-1] != obrasocial: continue # filtra si no pertenece a obra social pasada por parametro
				self.tabla.insert("", END, text="", values=(afiliado[0],afiliado[1], afiliado[2], afiliado[3],afiliado[4]))
			else:
				self.tabla.insert("", END, text="", values=(afiliado[0],afiliado[1], afiliado[2], afiliado[3],afiliado[4]))
		self.info_num_afiliados()
  
	def modo_agregar_afiliado(self):  
		self.modo="AGREGAR"
		self.abrir_ventana()
		
	def modo_editar_afiliado(self): #--- permite editar valores de la fila seleccion
		self.modo="EDITAR"  
		self.afiliado_target = self.get_datos_tabla() 
		if self.afiliado_target == [""]:
			messagebox.showinfo(message="POR FAVOR SELECCIONE ALGUN AFILIADO", title="EDITAR") # --- caja de mensaje
		else:
			self.abrir_ventana()
			self.activar_btns_toplevel()
			self.completar_campos_toplevel()

	def modo_eliminar_afiliado(self):
		self.modo="ELIMINAR"
		self.afiliado_target = self.get_datos_tabla() 
		if self.afiliado_target == [""]:
			messagebox.showinfo(message="POR FAVOR SELECCIONE ALGUN AFILIADO", title="ELIMINAR AFILIADO") # --- caja de mensaje
		else:
			if  messagebox.askyesno(message="ESTA ACCIÓN NO PODRÁ DESHACERSE\n\n            ¿DESEA CONTINUAR?", title=f"{(self.modo)} AFILIADO"): 
				self.eliminar_afiliado()
				
	def validar_nuevo_afiliado(self):
		self.desactivar_btns_toplevel()
		self.nuevo_afiliado = self.get_nuevo_afiliado()
		if self.afiliado_target == self.nuevo_afiliado:
			messagebox.showinfo(message="NO HA REALIZADO NINGUNA MODIFICACIÓN!!!", title="INFO")
			self.activar_btns_toplevel()
			return False
		if self.nuevo_afiliado[0]=="" or self.nuevo_afiliado[1]=="" or self.nuevo_afiliado[2]=="" or self.nuevo_afiliado[3]=="" or self.nuevo_afiliado[4]=="":
			messagebox.showinfo(message="TODOS LOS CAMPOS SON OBLIGATORIOS", title="INFO")
			self.activar_btns_toplevel()
			return False
		# --- valida campos numericos y alfabeticos
		valores_campo = [self.nuevo_afiliado[i].replace(" ","").isalpha() if i < 2 else self.nuevo_afiliado[i].replace(" ","").isnumeric() for i in range(4)]
		if valores_campo != [True,True,True,True]:
			messagebox.showinfo(message=" POR FAVOR RESPETE EL CARÁCTER ALFABÉTICO O NUMÉRICO \n\tDE LOS CAMPOS SEGÚN CORRESPONDA", title="INFO")
			self.activar_btns_toplevel()
			return False
		self.activar_btns_toplevel()
		return True

	def desactivar_btns_toplevel(self): # -- desactiva botones 
		self.btn_salir.configure(state='disabled')
		self.btn_ingresar.configure(state='disabled')

	def activar_btns_toplevel(self):
		self.btn_salir.configure(state='normal')
		self.btn_ingresar.configure(state='normal')

	def grabar_afiliado(self):
		self.nuevo_afiliado = self.get_nuevo_afiliado()
		if self.validar_nuevo_afiliado():  ########### si es verdadero
			self.desactivar_btns_toplevel()
			if self.modo=="AGREGAR":
				if messagebox.askyesno(message="ESTA ACCIÓN AGREGARÁ UN NUEVO AFILIADO...  \n\n\t   ¿DESEA CONTINUAR?", title=f"{(self.modo).upper()} AFILIADO"):
					messagebox.showinfo(message="EL AFILIADO HA SIDO AGREGADO CON ÉXITO", title="INFO")
					self.agregar_afiliado()
			else: # modo "EDITAR"
				if messagebox.askyesno(message="ESTA ACCIÓN MODIFICARÁ LOS DATOS DEL AFILIADO...  \n\n\t       ¿DESEA CONTINUAR?", title=f"{(self.modo).upper()} AFILIADO"):
					self.agregar_afiliado()
					self.eliminar_afiliado()
					messagebox.showinfo(message="LOS DATOS HAN SIDO MODIFICADOS CON ÉXITO", title="INFO")
			self.cerrar_toplevel()
			self.limpiar_tabla()
	
	def agregar_afiliado(self,archivo="afiliados/afiliados.csv"):
		try:
			with open(archivo, 'a', newline='') as csvfile:  
				writer_object = csv.writer(csvfile)
				writer_object.writerow(self.nuevo_afiliado)
		except Exception as e:
			messagebox.showerror(message=e, title="ERROR!!!")
		self.salir_ventana()

	def eliminar_afiliado(self,archivo="afiliados/afiliados.csv"):
		try:
			with open(archivo, 'r', encoding='latin1') as datos:
				nueva_lista_afiliados=[]
				csvreader = csv.reader(datos)
				for linea in csvreader:
					if linea == []: continue
					nueva_lista_afiliados.append(linea)	
		except Exception as e:
			messagebox.showerror(message=e, title="ERROR!!!")
		try:	
			with open(archivo, 'w', encoding='latin1') as datos:
				csvwriter = csv.writer(datos)
				for linea in nueva_lista_afiliados:
					if linea == self.afiliado_target: continue
					csvwriter.writerow(linea)
			if self.modo == "ELIMINAR":
				messagebox.showinfo(message="EL AFILIADO HA SIDO ELIMINADO CON ÉXITO", title="INFO")
		except Exception as e:
			messagebox.showerror(message=e, title="ERROR!!!")
		self.limpiar_tabla()

	def completar_campos_toplevel(self): # --- asigna valores a las variables de los campos (con valores de la fila seleccionada)
		self.apellido.set(self.afiliado_target[0]) 
		self.nombre.set(self.afiliado_target[1])
		self.dni_afiliado.set(self.afiliado_target[2])
		self.telefono.set(self.afiliado_target[3])
		self.obrasocial.set(self.afiliado_target[4])

	def limpiar_tabla(self):
		self.afiliados=[]
		self.eliminar_elementos_tabla()
		self.inicializar_tabla()

	def actualizar_tabla(self):
		self.eliminar_elementos_tabla()
		self.current_obrasocial = self.entry_obra_social.get() # --- toma valor asignado al campo obra social
		self.cargar_tabla(self.get_afiliados_sort(),self.get_obrasocial()) # --- carga tabla con los nuevos valores
		
	def eliminar_elementos_tabla(self):
		self.tabla.delete(*self.tabla.get_children()) # --- elimina todos los elementos de la tabla

	def ordenar_lista_afiliados(self):
		self.afiliados_sort= sorted(self.afiliados)

	def salir_ventana(self):
		self.cerrar_toplevel()
		self.ordenar_lista_afiliados()
		self.actualizar_tabla()

	def cerrar_toplevel(self):
		self.top_level.destroy()

	def get_afiliados_sort(self):
		return self.afiliados_sort

	def get_obrasocial(self):
		return self.current_obrasocial 

	def get_cantidad_afiliados(self):
		return len(self.tabla.get_children())

	def get_datos_tabla(self): # --- devuelve valores de la fila seleccionada en la tabla (treeview)
		focus_item = self.tabla.focus()
		datos_afil_target = self.tabla.item(focus_item)["values"]
		datos_afil_string= ",".join(str(elementos) for elementos in datos_afil_target)
		return datos_afil_string.split(",")

	def get_nuevo_afiliado(self):
		self.nuevo_afiliado = [
							self.apellido.get().upper(),
							self.nombre.get().upper(),
							self.dni_afiliado.get().upper(),
							self.telefono.get().upper(),
							self.obrasocial.get().upper()
							]
		return self.nuevo_afiliado

	def info_num_afiliados(self): # informa numero de afiliados en label de ventana principal
		self.text_info_frame1.set(f"NÚMERO DE AFILIADOS \t - {self.get_cantidad_afiliados()} -")

	def generar_documento(self,ruta="archivos_gen"): # genera archivo para imprimir
		if messagebox.askyesno(message="¿DESEA GENERAR UN DOCUMENTO CON ESTOS DATOS?", title=f"GENERAR DOCUMENTO"):
			c = 0
			doc = Document()
			sections=doc.sections
			for section in sections:
				section.top_margin = Cm(1.5)
				section.bottom_margin = Cm(1.5)
				section.left_margin = Cm(2)
				section.right_margin = Cm(2)
			style = doc.styles['Normal']
			style.paragraph_format.space_after = Pt(0) # Espacio entre lineas
			font = style.font
			font.name = 'Courier New'
			font.size = Pt(9)

			doc.add_paragraph('     APELLIDO               NOMBRE                     DNI/AFILIADO       TELEFONO     O.S.\n')
			for i in self.get_afiliados_sort():
				if self.get_obrasocial() != "TODAS":
					if i[4] == self.get_obrasocial(): 
						c+=1
						linea = f"{str(c).rjust(3,' ')}- {i[0][:22].ljust(22,' ')} {i[1][:22].ljust(22,' ')}    {i[2][:13].rjust(13,' ')}     {i[3][:10].rjust(10,' ')}     {i[4][:4].rjust(4,' ')}"
						doc.add_paragraph(linea)
				else: 
					c+=1
					linea = f"{str(c).rjust(3,' ')}- {i[0][:22].ljust(22,' ')} {i[1][:22].ljust(22,' ')}    {i[2][:13].rjust(13,' ')}     {i[3][:10].rjust(10,' ')}     {i[4][:4].rjust(4,' ')}"
					doc.add_paragraph(linea)

			fecha = datetime.now()
			nombre_archivo = f'{fecha.strftime("%Y-%m-%d_%H-%M-%S")}_{self.get_obrasocial()}'

			try:
				doc.save(f'{ruta}/{nombre_archivo}.docx')
				messagebox.showinfo(message=f"\n     EL ARCHIVO FUE GENERADO CON ÉXITO: {ruta}/{nombre_archivo}.docx\n", title="INFO")

			except Exception as e:
				messagebox.showerror(message=e, title="ERROR!!!")

	def abrir_carpeta_docs(self):
		ruta = os.getcwd() + "\\archivos_gen"
		archivo = filedialog.askopenfilename(title="DOCUMENTOS",initialdir=ruta, filetypes= [("Archivos word","*.docx")])
		try:
			os.system(archivo)
		except Exception as e:
			messagebox.showerror(message=e, title="ERROR!!!")

def app():
	v = App()
	v.root.mainloop()

if __name__ == "__main__":
	app()